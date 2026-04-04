from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.models import ActivityEvent, Anomaly, Incident, ReportSection

Locale = Literal["en", "ru"]


@dataclass(frozen=True)
class ReportMetric:
    label: str
    value: str
    detail: str


@dataclass(frozen=True)
class ReportTaskLine:
    title: str
    owner: str
    eta: str
    status: str
    notes: str


@dataclass(frozen=True)
class PreparedReport:
    locale: Locale
    file_stem: str
    title: str
    subtitle: str
    status_badge: str
    summary_title: str
    summary_body: str
    metrics: list[ReportMetric]
    facts_title: str
    facts: list[tuple[str, str]]
    sections: list[ReportSection]
    tasks_title: str
    tasks: list[ReportTaskLine]
    audit_title: str
    audit_lines: list[str]
    methodology_title: str
    methodology_points: list[str]


REGION_TRANSLATIONS = {
    "Atyrau Region": "Атырауская область",
    "Mangystau Region": "Мангистауская область",
    "Aktobe Region": "Актюбинская область",
    "West Kazakhstan Region": "Западно-Казахстанская область",
    "Kyzylorda Region": "Кызылординская область",
    "Pavlodar Region": "Павлодарская область",
    "Akmola Region": "Акмолинская область",
    "Almaty Region": "Алматинская область",
    "Almaty City": "Алматы",
    "Karaganda Region": "Карагандинская область",
    "Kostanay Region": "Костанайская область",
    "North Kazakhstan Region": "Северо-Казахстанская область",
    "East Kazakhstan Region": "Восточно-Казахстанская область",
    "Turkistan Region": "Туркестанская область",
    "Zhambyl Region": "Жамбылская область",
    "Ulytau Region": "область Улытау",
    "Kazakhstan": "Казахстан",
}

PLACE_TRANSLATIONS = {
    "Makat District": "Макатский район",
    "Zhanybek District": "Жанибекский район",
    "Martuk district": "Мартукский район",
    "Martok district": "Мартукский район",
    "Bayganin District": "Байганинский район",
    "Bayğanïn District": "Байганинский район",
    "Aqtöbe Region": "Актюбинская область",
    "Nura District": "Нуринский район",
    "Satbayev": "Сатпаев",
    "Tengiz Field": "месторождение Тенгиз",
}

ASSET_TRANSLATIONS = {
    "Tengiz satellite cluster": "Тенгизский спутниковый кластер",
    "Karabatan processing block": "Карабатанский перерабатывающий блок",
    "Mangystau export hub": "Мангистауский экспортный узел",
    "Aktobe compressor ring": "Актюбинское компрессорное кольцо",
    "Karachaganak gas train": "Карачаганакская газовая линия",
    "Kumkol gathering node": "Кумкольский узел сбора",
    "Pavlodar refinery corridor": "Павлодарский перерабатывающий коридор",
}

OWNER_TRANSLATIONS = {
    "Field integrity desk": "группа полевой целостности",
    "Ops coordinator": "координатор эксплуатации",
    "Reliability engineer": "инженер по надежности",
    "ESG lead": "руководитель ESG",
    "Response lead": "ответственный за реагирование",
    "Remote sensing analyst": "аналитик дистанционного зондирования",
    "Area operations coordinator": "координатор площадки",
    "Compliance lead": "руководитель по соблюдению требований",
    "ESG desk": "команда ESG",
    "MRV response lead": "координатор MRV-реагирования",
    "Earth Engine screening": "скрининг Earth Engine",
}

TASK_TITLE_TRANSLATIONS = {
    "Dispatch LDAR walkdown request": "Отправить запрос на LDAR-обход",
    "Cross-check flare line maintenance history": "Проверить историю обслуживания факельной линии",
    "Draft regulator-facing MRV note": "Подготовить MRV-заметку для регулятора",
    "Validate signal persistence against 12-week baseline": "Проверить устойчивость сигнала по базовому уровню за 12 недель",
    "Assign field verification owner": "Назначить ответственного за выездную проверку",
    "Collect operator comment": "Собрать комментарий оператора",
}

AUDIT_SOURCE_TRANSLATIONS = {
    "gee": "Google Earth Engine",
    "workflow": "рабочий процесс",
}

AUDIT_TITLE_TRANSLATIONS = {
    "Measurement evidence linked to incident": "Данные измерений привязаны к инциденту",
    "Incident created from screening signal": "Инцидент создан по итогам скрининга",
    "Verification task completed": "Задача проверки выполнена",
    "MRV report generated": "MRV-отчет сформирован",
}

EVIDENCE_SOURCE_TRANSLATIONS = {
    "Google Earth Engine / Sentinel-5P + VIIRS thermal context": "Google Earth Engine / Sentinel-5P + тепловой контекст VIIRS",
}

RECOMMENDED_ACTION_TRANSLATIONS = {
    "Keep this candidate in the manual review queue and confirm it with the next valid CH4 scene before escalation.": "Оставить эту зону в очереди на ручной разбор и подтвердить ее по следующей валидной сцене CH4 перед эскалацией.",
}


def prepare_report(
    anomaly: Anomaly,
    incident: Incident,
    audit_events: list[ActivityEvent],
    locale: Locale = "en",
) -> PreparedReport:
    labels = _labels(locale)
    completed_tasks = sum(1 for task in incident.tasks if task.status == "done")
    total_tasks = len(incident.tasks)
    thermal_hits = anomaly.night_thermal_hits_72h or anomaly.thermal_hits_72h or 0
    incident_status = _incident_status_label(incident.status, locale)
    screening_level = _severity_label(anomaly.severity, locale)

    metrics = [
        ReportMetric(
            label=labels["metric_priority"],
            value=incident.priority,
            detail=f"{labels['metric_owner']}: {_translate_owner(incident.owner, locale)}",
        ),
        ReportMetric(
            label=labels["metric_signal_score"],
            value=f"{anomaly.signal_score}/100",
            detail=screening_level,
        ),
        ReportMetric(
            label=labels["metric_methane_uplift"],
            value=_format_delta(anomaly),
            detail=_format_current_baseline(anomaly, locale),
        ),
        ReportMetric(
            label=labels["metric_thermal"],
            value=_format_thermal_hits(thermal_hits, locale),
            detail=labels["metric_thermal_detail"],
        ),
        ReportMetric(
            label=labels["metric_tasks"],
            value=f"{completed_tasks}/{total_tasks}",
            detail=_task_progress_detail(completed_tasks, total_tasks, locale),
        ),
        ReportMetric(
            label=labels["metric_window"],
            value=_translate_window(incident.verification_window, locale),
            detail=incident_status,
        ),
    ]

    facts = [
        (labels["generated"], incident.report_generated_at or labels["on_demand"]),
        (labels["incident"], incident.id),
        (labels["zone"], _translate_asset(anomaly.asset_name, locale)),
        (labels["region"], _translate_place(anomaly.region, locale)),
        (labels["coordinates"], anomaly.coordinates),
        (labels["verification_area"], _translate_place(anomaly.verification_area, locale)),
        (labels["nearest_address"], anomaly.nearest_address or labels["not_available"]),
        (labels["nearest_landmark"], _translate_place(anomaly.nearest_landmark, locale)),
        (labels["owner"], _translate_owner(incident.owner, locale)),
        (labels["priority"], incident.priority),
        (labels["status"], incident_status),
        (labels["window"], _translate_window(incident.verification_window, locale)),
        (labels["evidence_source"], _translate_evidence_source(anomaly.evidence_source, locale)),
        (labels["baseline_window"], _translate_baseline_window(anomaly.baseline_window, locale)),
    ]

    tasks = [
        ReportTaskLine(
            title=_translate_task_title(task.title, locale),
            owner=_translate_owner(task.owner, locale),
            eta=_format_hours(task.eta_hours, locale),
            status=labels["done"] if task.status == "done" else labels["open"],
            notes=task.notes.strip(),
        )
        for task in incident.tasks
    ]

    return PreparedReport(
        locale=locale,
        file_stem=f"{incident.id.lower()}-mrv-report",
        title=f"{labels['title']}: {incident.id}",
        subtitle=labels["subtitle"],
        status_badge=f"{incident.priority} · {incident_status}",
        summary_title=labels["summary_title"],
        summary_body=_summary_text(anomaly, incident, completed_tasks, locale),
        metrics=metrics,
        facts_title=labels["facts_title"],
        facts=facts,
        sections=_localized_sections(anomaly, incident, completed_tasks, locale),
        tasks_title=labels["tasks_title"],
        tasks=tasks,
        audit_title=labels["audit_title"],
        audit_lines=[_format_audit_line(event, locale) for event in audit_events],
        methodology_title=labels["methodology_title"],
        methodology_points=_methodology_points(locale),
    )


def render_html(report: PreparedReport, auto_print: bool = False) -> str:
    labels = _labels(report.locale)
    metrics_markup = "".join(
        (
            "<div class='metric-card'>"
            f"<span class='metric-label'>{_escape_html(metric.label)}</span>"
            f"<strong class='metric-value'>{_escape_html(metric.value)}</strong>"
            f"<span class='metric-detail'>{_escape_html(metric.detail)}</span>"
            "</div>"
        )
        for metric in report.metrics
    )
    facts_markup = "".join(
        (
            "<tr>"
            f"<th>{_escape_html(label)}</th>"
            f"<td>{_escape_html(value)}</td>"
            "</tr>"
        )
        for label, value in report.facts
    )
    section_markup = "".join(
        (
            "<section class='report-section'>"
            f"<h2>{_escape_html(section.title)}</h2>"
            f"<p>{_escape_html(section.body)}</p>"
            "</section>"
        )
        for section in report.sections
    )
    task_rows = "".join(
        (
            "<tr>"
            f"<td>{_escape_html(task.title)}</td>"
            f"<td>{_escape_html(task.owner)}</td>"
            f"<td>{_escape_html(task.eta)}</td>"
            f"<td>{_escape_html(task.status)}</td>"
            f"<td>{_escape_html(task.notes or '—')}</td>"
            "</tr>"
        )
        for task in report.tasks
    )
    audit_rows = "".join(f"<li>{_escape_html(line)}</li>" for line in report.audit_lines)
    methodology_rows = "".join(
        f"<li>{_escape_html(point)}</li>" for point in report.methodology_points
    )
    auto_print_script = (
        "<script>window.addEventListener('load',()=>{setTimeout(()=>window.print(),120);});</script>"
        if auto_print
        else ""
    )

    return (
        "<!doctype html>"
        f"<html lang='{report.locale}'>"
        "<head>"
        "<meta charset='utf-8' />"
        "<meta name='viewport' content='width=device-width, initial-scale=1' />"
        f"<title>{_escape_html(report.title)}</title>"
        "<style>"
        "body{margin:0;font-family:'Segoe UI',Arial,sans-serif;background:#f4eee6;color:#122331;}"
        "main{max-width:1040px;margin:24px auto;padding:28px;border:1px solid #d8d0c6;background:#fffdfa;}"
        ".header{display:flex;justify-content:space-between;gap:20px;align-items:flex-start;border-bottom:1px solid #e4ddd3;padding-bottom:20px;}"
        ".eyebrow{display:inline-block;padding:6px 10px;border-radius:999px;background:#163244;color:#fff;font-size:12px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;}"
        "h1{margin:14px 0 6px;font-size:32px;line-height:1.15;}"
        ".subtitle{margin:0;color:#5f6e78;line-height:1.6;max-width:720px;}"
        ".toolbar button{padding:10px 14px;border:1px solid #c8d3d9;background:#fff;font:inherit;cursor:pointer;}"
        ".summary{margin-top:22px;padding:18px 20px;border:1px solid #d7e2e7;background:#f5fafb;}"
        ".summary h2{margin:0 0 8px;font-size:16px;}"
        ".summary p{margin:0;line-height:1.65;}"
        ".metrics{margin-top:20px;display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:14px;}"
        ".metric-card{padding:14px 16px;border:1px solid #d9e1e7;background:#fff;min-height:104px;}"
        ".metric-label{display:block;font-size:12px;letter-spacing:.08em;text-transform:uppercase;color:#657784;}"
        ".metric-value{display:block;margin-top:8px;font-size:24px;line-height:1.15;color:#132633;}"
        ".metric-detail{display:block;margin-top:8px;font-size:13px;line-height:1.5;color:#5f6e78;}"
        ".facts{margin-top:24px;}"
        ".facts h2,.report-section h2,.table-block h2,.audit h2,.methodology h2{margin:0 0 12px;font-size:18px;}"
        "table{width:100%;border-collapse:collapse;}"
        ".facts-table th,.facts-table td{padding:10px 12px;border:1px solid #d8e0e5;text-align:left;vertical-align:top;}"
        ".facts-table th{width:30%;background:#f7fafc;font-size:12px;letter-spacing:.08em;text-transform:uppercase;color:#5b6d79;}"
        ".report-section{margin-top:20px;}"
        ".report-section p{margin:0;line-height:1.7;}"
        ".table-block{margin-top:22px;}"
        ".tasks-table th,.tasks-table td{padding:10px 12px;border:1px solid #d8e0e5;text-align:left;vertical-align:top;}"
        ".tasks-table th{background:#f7fafc;font-size:12px;letter-spacing:.08em;text-transform:uppercase;color:#5b6d79;}"
        ".audit,.methodology{margin-top:22px;}"
        ".audit ul,.methodology ul{margin:0;padding-left:20px;line-height:1.7;}"
        "@media print{body{background:#fff;}main{margin:0;border:none;padding:0;} .toolbar{display:none;} .metrics{gap:10px;} .metric-card{min-height:auto;}}"
        "</style>"
        "</head>"
        "<body>"
        "<main>"
        "<div class='header'>"
        "<div>"
        f"<span class='eyebrow'>{_escape_html(report.status_badge)}</span>"
        f"<h1>{_escape_html(report.title)}</h1>"
        f"<p class='subtitle'>{_escape_html(report.subtitle)}</p>"
        "</div>"
        "<div class='toolbar'>"
        f"<button onclick='window.print()'>{_escape_html(labels['print'])}</button>"
        "</div>"
        "</div>"
        "<section class='summary'>"
        f"<h2>{_escape_html(report.summary_title)}</h2>"
        f"<p>{_escape_html(report.summary_body)}</p>"
        "</section>"
        f"<section class='metrics'>{metrics_markup}</section>"
        "<section class='facts'>"
        f"<h2>{_escape_html(report.facts_title)}</h2>"
        f"<table class='facts-table'>{facts_markup}</table>"
        "</section>"
        f"{section_markup}"
        "<section class='table-block'>"
        f"<h2>{_escape_html(report.tasks_title)}</h2>"
        "<table class='tasks-table'>"
        "<thead><tr>"
        f"<th>{_escape_html(labels['task_col_title'])}</th>"
        f"<th>{_escape_html(labels['task_col_owner'])}</th>"
        f"<th>{_escape_html(labels['task_col_eta'])}</th>"
        f"<th>{_escape_html(labels['task_col_status'])}</th>"
        f"<th>{_escape_html(labels['task_col_notes'])}</th>"
        "</tr></thead>"
        f"<tbody>{task_rows}</tbody>"
        "</table>"
        "</section>"
        "<section class='audit'>"
        f"<h2>{_escape_html(report.audit_title)}</h2>"
        f"<ul>{audit_rows}</ul>"
        "</section>"
        "<section class='methodology'>"
        f"<h2>{_escape_html(report.methodology_title)}</h2>"
        f"<ul>{methodology_rows}</ul>"
        "</section>"
        f"{auto_print_script}"
        "</main>"
        "</body></html>"
    )


def render_pdf(report: PreparedReport) -> bytes:
    regular_font, bold_font = _resolve_pdf_fonts()
    labels = _labels(report.locale)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontName=bold_font,
        fontSize=22,
        leading=26,
        textColor=colors.HexColor("#10212b"),
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#5e6d77"),
        spaceAfter=12,
    )
    badge_style = ParagraphStyle(
        "ReportBadge",
        parent=styles["BodyText"],
        fontName=bold_font,
        fontSize=9,
        leading=12,
        textColor=colors.white,
        backColor=colors.HexColor("#163244"),
        borderPadding=(4, 8, 4, 8),
        spaceAfter=10,
    )
    section_style = ParagraphStyle(
        "ReportSection",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#132633"),
        spaceBefore=14,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#33434f"),
    )
    small_style = ParagraphStyle(
        "ReportSmall",
        parent=body_style,
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#5f6e78"),
    )

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )

    elements: list[object] = [
        Paragraph(_escape_html(report.status_badge), badge_style),
        Paragraph(_escape_html(report.title), title_style),
        Paragraph(_escape_html(report.subtitle), subtitle_style),
    ]

    summary_box = Table(
        [[Paragraph(f"<b>{_escape_html(report.summary_title)}</b><br/>{_escape_html(report.summary_body)}", body_style)]],
        colWidths=[178 * mm],
        hAlign="LEFT",
    )
    summary_box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f5fafb")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d8e4e8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    elements.extend([summary_box, Spacer(1, 10)])

    metric_rows: list[list[object]] = []
    for index in range(0, len(report.metrics), 3):
        cards = report.metrics[index:index + 3]
        while len(cards) < 3:
            cards.append(ReportMetric("", "", ""))
        metric_rows.append(
            [
                Paragraph(
                    (
                        f"<font size='8'><b>{_escape_html(card.label)}</b></font><br/>"
                        f"<font size='16'><b>{_escape_html(card.value)}</b></font><br/>"
                        f"<font size='8'>{_escape_html(card.detail)}</font>"
                    )
                    if card.label
                    else "",
                    body_style,
                )
                for card in cards
            ]
        )
    metric_table = Table(metric_rows, colWidths=[58 * mm, 58 * mm, 58 * mm], hAlign="LEFT")
    metric_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#dae3e7")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#dae3e7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    elements.extend([metric_table, Spacer(1, 10)])

    elements.append(Paragraph(_escape_html(report.facts_title), section_style))
    fact_rows = [
        [
            Paragraph(f"<b>{_escape_html(label)}</b>", body_style),
            Paragraph(_escape_html(value), body_style),
        ]
        for label, value in report.facts
    ]
    facts_table = Table(fact_rows, colWidths=[48 * mm, 130 * mm], hAlign="LEFT")
    facts_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f7fafc")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#d8e0e5")),
                ("INNERGRID", (0, 0), (-1, -1), 0.55, colors.HexColor("#d8e0e5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    elements.extend([facts_table, Spacer(1, 10)])

    for section in report.sections:
        elements.append(Paragraph(_escape_html(section.title), section_style))
        elements.append(Paragraph(_escape_html(section.body), body_style))

    elements.append(Paragraph(_escape_html(report.tasks_title), section_style))
    task_rows = [
        [
            Paragraph(f"<b>{_escape_html(labels['task_col_title'])}</b>", small_style),
            Paragraph(f"<b>{_escape_html(labels['task_col_owner'])}</b>", small_style),
            Paragraph(f"<b>{_escape_html(labels['task_col_eta'])}</b>", small_style),
            Paragraph(f"<b>{_escape_html(labels['task_col_status'])}</b>", small_style),
            Paragraph(f"<b>{_escape_html(labels['task_col_notes'])}</b>", small_style),
        ]
    ]
    for task in report.tasks:
        task_rows.append(
            [
                Paragraph(_escape_html(task.title), body_style),
                Paragraph(_escape_html(task.owner), body_style),
                Paragraph(_escape_html(task.eta), body_style),
                Paragraph(_escape_html(task.status), body_style),
                Paragraph(_escape_html(task.notes or "—"), body_style),
            ]
        )
    tasks_table = Table(
        task_rows,
        colWidths=[58 * mm, 38 * mm, 18 * mm, 24 * mm, 40 * mm],
        hAlign="LEFT",
        repeatRows=1,
    )
    tasks_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f7fafc")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#d8e0e5")),
                ("INNERGRID", (0, 0), (-1, -1), 0.55, colors.HexColor("#d8e0e5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.extend([tasks_table, Spacer(1, 8)])

    elements.append(Paragraph(_escape_html(report.audit_title), section_style))
    for line in report.audit_lines:
        elements.append(Paragraph(f"• {_escape_html(line)}", body_style))
        elements.append(Spacer(1, 3))

    elements.append(Paragraph(_escape_html(report.methodology_title), section_style))
    for point in report.methodology_points:
        elements.append(Paragraph(f"• {_escape_html(point)}", body_style))
        elements.append(Spacer(1, 3))

    document.build(elements)
    return buffer.getvalue()


def render_docx(report: PreparedReport) -> bytes:
    labels = _labels(report.locale)
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Arial"
    heading_1.font.size = Pt(20)
    heading_1.font.color.rgb = RGBColor(0x10, 0x21, 0x2B)

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Arial"
    heading_2.font.size = Pt(13)
    heading_2.font.color.rgb = RGBColor(0x13, 0x26, 0x33)

    document.add_paragraph(report.status_badge)
    title = document.add_heading(report.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(report.subtitle)

    document.add_heading(report.summary_title, level=2)
    document.add_paragraph(report.summary_body)

    document.add_heading("Ключевые показатели" if report.locale == "ru" else "Key metrics", level=2)
    metric_table = document.add_table(rows=1, cols=3)
    metric_table.style = "Table Grid"
    header_cells = metric_table.rows[0].cells
    for index, metric in enumerate(report.metrics[:3]):
        header_cells[index].text = f"{metric.label}\n{metric.value}\n{metric.detail}"
    for row_start in range(3, len(report.metrics), 3):
        row = metric_table.add_row().cells
        chunk = report.metrics[row_start:row_start + 3]
        for index, metric in enumerate(chunk):
            row[index].text = f"{metric.label}\n{metric.value}\n{metric.detail}"

    document.add_heading(report.facts_title, level=2)
    facts_table = document.add_table(rows=0, cols=2)
    facts_table.style = "Table Grid"
    for label, value in report.facts:
        row = facts_table.add_row().cells
        row[0].text = label
        row[1].text = value

    for section in report.sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(section.body)

    document.add_heading(report.tasks_title, level=2)
    tasks_table = document.add_table(rows=1, cols=5)
    tasks_table.style = "Table Grid"
    task_headers = [
        labels["task_col_title"],
        labels["task_col_owner"],
        labels["task_col_eta"],
        labels["task_col_status"],
        labels["task_col_notes"],
    ]
    for index, header in enumerate(task_headers):
        tasks_table.rows[0].cells[index].text = header
    for task in report.tasks:
        row = tasks_table.add_row().cells
        row[0].text = task.title
        row[1].text = task.owner
        row[2].text = task.eta
        row[3].text = task.status
        row[4].text = task.notes or "—"

    document.add_heading(report.audit_title, level=2)
    for line in report.audit_lines:
        document.add_paragraph(line, style="List Bullet")

    document.add_heading(report.methodology_title, level=2)
    for point in report.methodology_points:
        document.add_paragraph(point, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _labels(locale: Locale) -> dict[str, str]:
    if locale == "ru":
        return {
            "title": "Отчет Saryna MRV",
            "subtitle": "Рабочий отчет по кейсу: зона, ход проверки, ответственные и журнал действий.",
            "summary_title": "Итог по кейсу",
            "facts_title": "Карточка кейса",
            "tasks_title": "Задачи на проверку",
            "audit_title": "Журнал действий",
            "methodology_title": "Основа отчета",
            "generated": "Сформирован",
            "on_demand": "По запросу",
            "incident": "Инцидент",
            "zone": "Подозрительная зона",
            "region": "Регион",
            "coordinates": "Координаты",
            "verification_area": "Район проверки",
            "nearest_address": "Ближайший адрес",
            "nearest_landmark": "Ближайший ориентир",
            "owner": "Ответственный",
            "priority": "Приоритет",
            "status": "Статус кейса",
            "window": "Срок проверки",
            "evidence_source": "Источник данных",
            "baseline_window": "Базовое окно сравнения",
            "metric_priority": "Приоритет кейса",
            "metric_owner": "Ответственный",
            "metric_signal_score": "Оценка зоны",
            "metric_methane_uplift": "Рост CH4",
            "metric_thermal": "Тепловой контекст",
            "metric_thermal_detail": "Количество ночных срабатываний VIIRS за 72 часа",
            "metric_tasks": "Выполнено задач",
            "metric_window": "Окно проверки",
            "task_col_title": "Задача",
            "task_col_owner": "Ответственный",
            "task_col_eta": "Срок",
            "task_col_status": "Статус",
            "task_col_notes": "Примечание",
            "done": "Выполнено",
            "open": "Открыто",
            "not_available": "Нет данных",
            "print": "Печать / Сохранить как PDF",
        }

    return {
        "title": "Saryna MRV Report",
        "subtitle": "Operational report for the case: suspected zone, verification progress, owners, and audit trail.",
        "summary_title": "Case summary",
        "facts_title": "Case facts",
        "tasks_title": "Verification tasks",
        "audit_title": "Audit Timeline",
        "methodology_title": "Report basis",
        "generated": "Generated",
        "on_demand": "On demand",
        "incident": "Incident",
        "zone": "Suspected zone",
        "region": "Region",
        "coordinates": "Coordinates",
        "verification_area": "Verification area",
        "nearest_address": "Nearest address",
        "nearest_landmark": "Nearest landmark",
        "owner": "Owner",
        "priority": "Priority",
        "status": "Case status",
        "window": "Verification window",
        "evidence_source": "Evidence source",
        "baseline_window": "Baseline comparison window",
        "metric_priority": "Case priority",
        "metric_owner": "Owner",
        "metric_signal_score": "Zone score",
        "metric_methane_uplift": "CH4 uplift",
        "metric_thermal": "Thermal context",
        "metric_thermal_detail": "Night-time VIIRS detections over 72 hours",
        "metric_tasks": "Tasks completed",
        "metric_window": "Response window",
        "task_col_title": "Task",
        "task_col_owner": "Owner",
        "task_col_eta": "ETA",
        "task_col_status": "Status",
        "task_col_notes": "Notes",
        "done": "Done",
        "open": "Open",
        "not_available": "Not available",
        "print": "Print / Save as PDF",
    }


def _localized_sections(
    anomaly: Anomaly,
    incident: Incident,
    completed_tasks: int,
    locale: Locale,
) -> list[ReportSection]:
    screening_title = "Что показали данные" if locale == "ru" else "What the data showed"
    operations_title = "Состояние кейса" if locale == "ru" else "Current case status"
    next_step_title = "Следующее действие" if locale == "ru" else "Recommended next step"
    return [
        ReportSection(
            title=screening_title,
            body=_screening_section_text(anomaly, locale),
        ),
        ReportSection(
            title=operations_title,
            body=_operations_section_text(anomaly, incident, completed_tasks, locale),
        ),
        ReportSection(
            title=next_step_title,
            body=_next_step_section_text(anomaly, incident, locale),
        ),
    ]


def _summary_text(anomaly: Anomaly, incident: Incident, completed_tasks: int, locale: Locale) -> str:
    zone = _translate_asset(anomaly.asset_name, locale)
    region = _translate_place(anomaly.region, locale)
    thermal_hits = anomaly.night_thermal_hits_72h or anomaly.thermal_hits_72h or 0
    owner = _translate_owner(incident.owner, locale)
    uplift = _format_delta(anomaly)
    if locale == "ru":
        thermal_text = (
            f"За последние 72 часа рядом зафиксировано {thermal_hits} ночных срабатываний VIIRS."
            if thermal_hits
            else "За последние 72 часа ночные срабатывания VIIRS рядом не зафиксированы."
        )
        return (
            f"Кейс {incident.id} открыт по зоне «{zone}» в регионе «{region}». "
            f"Последняя спутниковая сцена показала рост метана {uplift}. "
            f"{thermal_text} "
            f"Ответственный за кейс: {owner}. Выполнено задач: {completed_tasks} из {len(incident.tasks)}."
        )
    thermal_text = (
        f"{thermal_hits} night-time VIIRS detections were recorded nearby over the last 72 hours."
        if thermal_hits
        else "No night-time VIIRS detections were recorded nearby over the last 72 hours."
    )
    return (
        f"Case {incident.id} was opened for '{zone}' in {region}. "
        f"The latest satellite scene showed CH4 uplift of {uplift}. "
        f"{thermal_text} "
        f"Case owner: {owner}. Completed tasks: {completed_tasks} of {len(incident.tasks)}."
    )


def _screening_section_text(anomaly: Anomaly, locale: Locale) -> str:
    zone = _translate_asset(anomaly.asset_name, locale)
    region = _translate_place(anomaly.region, locale)
    verification_area = _translate_place(anomaly.verification_area, locale)
    landmark = _translate_place(anomaly.nearest_landmark, locale)
    address = anomaly.nearest_address
    if locale == "ru":
        text = (
            f"Зона «{zone}» отобрана после сравнения последней сцены с базовым окном: "
            f"{_format_delta(anomaly)} при уровне {anomaly.current_ch4_ppb:.2f} ppb против базового "
            f"{anomaly.baseline_ch4_ppb:.2f} ppb. Регион: {region}."
        )
        if verification_area != _labels(locale)["not_available"]:
            text += f" Район проверки: {verification_area}."
        if landmark != _labels(locale)["not_available"]:
            text += f" Ближайший ориентир: {landmark}."
        if address:
            text += f" Ближайший адрес: {address}."
        return text
    text = (
        f"Zone '{zone}' was selected after comparing the latest scene against the baseline window: "
        f"{_format_delta(anomaly)} with {anomaly.current_ch4_ppb:.2f} ppb against "
        f"{anomaly.baseline_ch4_ppb:.2f} ppb baseline. Region: {region}."
    )
    if verification_area != _labels(locale)["not_available"]:
        text += f" Verification area: {verification_area}."
    if landmark != _labels(locale)["not_available"]:
        text += f" Nearest landmark: {landmark}."
    if address:
        text += f" Nearest address: {address}."
    return text


def _operations_section_text(
    anomaly: Anomaly,
    incident: Incident,
    completed_tasks: int,
    locale: Locale,
) -> str:
    thermal_hits = anomaly.night_thermal_hits_72h or anomaly.thermal_hits_72h or 0
    owner = _translate_owner(incident.owner, locale)
    status_label = _incident_status_label(incident.status, locale)
    if locale == "ru":
        return (
            f"Кейс находится в статусе «{status_label}» и ведется с приоритетом {incident.priority}. "
            f"Ответственный: {owner}. Выполнено {completed_tasks} из {len(incident.tasks)} задач. "
            f"Тепловой контекст за 72 часа: {_format_thermal_hits(thermal_hits, locale)}."
        )
    return (
        f"The case is currently in '{status_label}' status with {incident.priority} priority. "
        f"Owner: {owner}. {completed_tasks} of {len(incident.tasks)} tasks are complete. "
        f"Thermal context over 72 hours: {_format_thermal_hits(thermal_hits, locale)}."
    )


def _next_step_section_text(anomaly: Anomaly, incident: Incident, locale: Locale) -> str:
    recommended_action = anomaly.recommended_action.strip() or (
        "Continue verification tasks and update the MRV record."
        if locale == "en"
        else "Продолжить проверку и обновить MRV-запись."
    )
    recommended_action = _translate_recommended_action(recommended_action, locale)
    if locale == "ru":
        return (
            f"Рекомендуемое действие по кейсу: {recommended_action} "
            f"Окно проверки: {_translate_window(incident.verification_window, locale)}."
        )
    return (
        f"Recommended next step: {recommended_action} "
        f"Verification window: {_translate_window(incident.verification_window, locale)}."
    )


def _methodology_points(locale: Locale) -> list[str]:
    if locale == "ru":
        return [
            "Слой CH4 построен на данных Sentinel-5P / TROPOMI и сравнивается с предыдущим базовым окном по Казахстану.",
            "Тепловой контекст берется из VIIRS Nightfire за последние 72 часа вокруг выбранной зоны.",
            "Отчет объединяет спутниковые показатели, состояние задач и журнал действий по кейсу.",
        ]
    return [
        "The CH4 layer is built from Sentinel-5P / TROPOMI and compared against the previous Kazakhstan baseline window.",
        "Thermal context comes from VIIRS Nightfire over the last 72 hours around the selected zone.",
        "The report combines satellite evidence, task progress, and the audit trail for the case.",
    ]


def _incident_status_label(value: str, locale: Locale) -> str:
    if locale == "ru":
        return {
            "triage": "первичный разбор",
            "verification": "проверка",
            "mitigation": "меры реагирования",
        }.get(value, value)
    return value


def _severity_label(value: str, locale: Locale) -> str:
    if locale == "ru":
        return {
            "high": "высокий приоритет",
            "medium": "средний приоритет",
            "watch": "режим наблюдения",
        }.get(value, value)
    return value


def _format_delta(anomaly: Anomaly) -> str:
    if anomaly.methane_delta_ppb is None:
        return f"{anomaly.methane_delta_pct:.2f}%"
    return f"{anomaly.methane_delta_ppb:.2f} ppb / {anomaly.methane_delta_pct:.2f}%"


def _format_current_baseline(anomaly: Anomaly, locale: Locale) -> str:
    if anomaly.current_ch4_ppb is None or anomaly.baseline_ch4_ppb is None:
        return _labels(locale)["not_available"]
    return f"{anomaly.current_ch4_ppb:.2f} ppb vs {anomaly.baseline_ch4_ppb:.2f} ppb"


def _format_thermal_hits(count: int, locale: Locale) -> str:
    if locale == "ru":
        return f"{count} срабатываний" if count != 1 else "1 срабатывание"
    return f"{count} detections" if count != 1 else "1 detection"


def _task_progress_detail(completed: int, total: int, locale: Locale) -> str:
    if locale == "ru":
        return f"Выполнено {completed} из {total} задач"
    return f"{completed} of {total} tasks completed"


def _format_audit_line(event: ActivityEvent, locale: Locale) -> str:
    source = AUDIT_SOURCE_TRANSLATIONS.get(event.source, event.source) if locale == "ru" else event.source
    if locale == "ru":
        title = AUDIT_TITLE_TRANSLATIONS.get(event.title, event.title)
        detail = _translate_audit_detail(event.detail, locale)
        actor = _translate_owner(event.actor, locale)
        return f"{event.occurred_at} | {title} | {detail} | источник: {source} | исполнитель: {actor}"
    return f"{event.occurred_at} | {event.title} | {event.detail} | source: {source} | actor: {event.actor}"


def _translate_place(value: str | None, locale: Locale) -> str:
    if not value:
        return _labels(locale)["not_available"]
    if locale != "ru":
        return value
    return _replace_known_place_parts(value)


def _translate_asset(value: str, locale: Locale) -> str:
    if locale != "ru":
        return value
    return ASSET_TRANSLATIONS.get(value, _translate_place(value, locale))


def _translate_owner(value: str, locale: Locale) -> str:
    if not value:
        return _labels(locale)["not_available"]
    return OWNER_TRANSLATIONS.get(value, value) if locale == "ru" else value


def _translate_task_title(value: str, locale: Locale) -> str:
    if locale != "ru":
        return value
    return TASK_TITLE_TRANSLATIONS.get(value, value)


def _translate_window(value: str, locale: Locale) -> str:
    if locale == "ru":
        return {
            "Next 12 hours": "в ближайшие 12 часов",
            "Next 24 hours": "в ближайшие 24 часа",
            "Next 48 hours": "в ближайшие 48 часов",
        }.get(value, value)
    return value


def _replace_known_place_parts(value: str) -> str:
    translated = value
    merged = {**PLACE_TRANSLATIONS, **REGION_TRANSLATIONS, **ASSET_TRANSLATIONS}
    for source, target in sorted(merged.items(), key=lambda item: len(item[0]), reverse=True):
        translated = translated.replace(source, target)
    return translated


def _translate_evidence_source(value: str | None, locale: Locale) -> str:
    if not value:
        return _labels(locale)["not_available"]
    if locale != "ru":
        return value
    return EVIDENCE_SOURCE_TRANSLATIONS.get(value, _replace_known_place_parts(value))


def _translate_baseline_window(value: str | None, locale: Locale) -> str:
    if not value:
        return _labels(locale)["not_available"]
    if locale != "ru":
        return value
    match = re.match(
        r"(?P<days>\d+)-day Kazakhstan baseline before (?P<timestamp>.+?); (?P<count>\d+) recent valid scenes checked\.",
        value,
    )
    if match:
        days = match.group("days")
        timestamp = match.group("timestamp")
        count = match.group("count")
        return f"{days}-дневный базовый уровень по Казахстану до {timestamp}; проверено {count} последних валидных сцен."
    return _replace_known_place_parts(value)


def _translate_recommended_action(value: str, locale: Locale) -> str:
    if locale != "ru":
        return value
    return RECOMMENDED_ACTION_TRANSLATIONS.get(value, _replace_known_place_parts(value))


def _translate_audit_detail(value: str, locale: Locale) -> str:
    if locale != "ru":
        return value
    translated = _replace_known_place_parts(value)
    patterns: list[tuple[re.Pattern[str], str]] = [
        (
            re.compile(r"^(INC-\d{8}-\d{2}) now has an updated MRV summary for stakeholder review\.$"),
            r"\1 получил обновленный MRV-отчет для внутреннего разбора.",
        ),
        (
            re.compile(r"^(INC-\d{8}-\d{2}-TASK-\d+) was marked done for (INC-\d{8}-\d{2})\.$"),
            r"\1 отмечена как выполненная по кейсу \2.",
        ),
        (
            re.compile(r"^(.+?) was promoted into (INC-\d{8}-\d{2}) with owner (.+?)\.$"),
            r"Зона «\1» переведена в \2; ответственный: \3.",
        ),
        (
            re.compile(r"^(GEE-\d{8}-\d{2}) screening evidence for (.+?) was attached to (INC-\d{8}-\d{2}) before escalation\.$"),
            r"Данные скрининга \1 по зоне «\2» прикреплены к \3 перед эскалацией.",
        ),
    ]
    for pattern, replacement in patterns:
        if pattern.match(translated):
            translated = pattern.sub(replacement, translated)
            break
    translated = _translate_owner(translated, locale)
    return translated


def _format_hours(hours: int, locale: Locale) -> str:
    return f"{hours} ч" if locale == "ru" else f"{hours}h"


def _escape_html(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def _resolve_pdf_fonts() -> tuple[str, str]:
    font_candidates = [
        (
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            "DejaVuSans",
            "DejaVuSansBold",
        ),
        (
            Path("C:/Windows/Fonts/DejaVuSans.ttf"),
            Path("C:/Windows/Fonts/DejaVuSans-Bold.ttf"),
            "DejaVuSans",
            "DejaVuSansBold",
        ),
        (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            "ArialUnicode",
            "ArialUnicodeBold",
        ),
    ]

    for regular_path, bold_path, regular_name, bold_name in font_candidates:
        if regular_path.exists() and bold_path.exists():
            if regular_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(regular_name, str(regular_path)))
            if bold_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(bold_name, str(bold_path)))
            return regular_name, bold_name

    return "Helvetica", "Helvetica-Bold"
